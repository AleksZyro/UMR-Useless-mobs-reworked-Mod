from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Mobliste.docx"
SECTION_TITLE = "Projektstand 24.08.2026"


ROWS = [
    ("Korrumpierter Silberfisch", "AKTIV / REFERENZ", "Aktives V5-Spielmodell, Textur, Animation, Sounds und Multipart-Hitbox sind eingebunden und durch den Projektprüfer validiert.", "Weitere Ingame-Politur nur nach Vergleichsbild."),
    ("Frost Stray", "NEU GENERIERT", "Komplett neues Tripo-v3.1-Modell mit 4K-Textur; exakter Mesh-Export mit 714'146 Dreiecken und sechs animierbaren Körperregionen.", "Abschliessende isolierte Ingame-Sichtprüfung von Grösse und Bodenhöhe."),
    ("Coral Drowned", "NEU GENERIERT", "Komplett neues Tripo-v3.1-Modell mit 4K-Textur; exakter Mesh-Export mit 728'013 Dreiecken und sechs animierbaren Körperregionen.", "Abschliessende isolierte Ingame-Sichtprüfung von Vorderseite, Grösse und Bodenhöhe."),
    ("Living Axolotl", "NEU GENERIERT", "Komplett neues langes Axolotl-Modell mit symmetrischen Kiemen und 4K-Textur; 733'834 Dreiecke und sieben animierbare Körperregionen.", "Abschliessende isolierte Ingame-Sichtprüfung von Körperachse, Grösse und Bodenhöhe."),
    ("Living Polar Bear", "BEREINIGT", "Eigenes Entity-, Spawn-Ei-, Renderer-, Mesh- und Texturpaket ist vorhanden; die abgetrennte 49'098-Dreieck-Fremdkomponente wird nach Zusammenhangsprüfung nicht mehr exportiert.", "Isolierter Ingame-Abschlusstest von Kopf, Bodenhöhe und Laufanimation."),
    ("Helping Allay", "EINGEBUNDEN / QA", "Eigenes Entity, Spawn-Ei und Exact-Layer sind vorhanden.", "Visuelle Ingame-Prüfung und gegebenenfalls Achsen-/Rig-Korrektur."),
    ("Oktopus", "EINGEBUNDEN / QA", "Eigenes Entity, Renderer, Modell, Textur, Wasser-KI und Sounds sind vorhanden.", "Isolation in der Modellgalerie und natürliche Tentakelanimation prüfen."),
    ("Squid / Kalmar", "EINGEBUNDEN / QA", "Eigene Living-Squid-Variante mit Spawn-Ei, Renderer, Mesh und Textur ist vorhanden.", "Grösse, acht Arme plus zwei Fangtentakel und Schwimmanimation prüfen."),
    ("Glow Squid", "EINGEBUNDEN / QA", "Eigene Living-Glow-Squid-Variante mit Spawn-Ei, Renderer, Mesh und Textur ist vorhanden.", "Blend-/Lichteffekt und Ingame-Grösse prüfen."),
    ("Living Ocelot", "EINGEBUNDEN / QA", "Eigene Entity, Spawn-Ei, Renderer, Mesh und Textur sind vorhanden.", "Bodenkontakt, Laufrichtung und Rekrutierungsverhalten prüfen."),
    ("Living Bat", "EINGEBUNDEN / QA", "Eigene Entity, Spawn-Ei, Renderer, Mesh, Textur und Fähigkeiten sind vorhanden.", "Fluganimation und Kollisionsgrösse prüfen."),
    ("Rooted Husk", "EINGEBUNDEN / QA", "Eigene Entity, Spawn-Ei, Renderer, Mesh, Textur und Fähigkeiten sind vorhanden.", "Bodenhöhe, Laufrichtung und Modellgrösse prüfen."),
    ("Web Cave Spider", "EINGEBUNDEN / QA", "Eigene Entity, Renderer, Mesh, Textur, Netzangriff und Sounds sind vorhanden.", "Bodenkontakt, Laufrichtung und Beinanimation prüfen."),
    ("Living Boss", "EINGEBUNDEN / QA", "Boss-Entity, Renderer, Tripo-Mesh, 4K-Textur, Phasen, Spezialangriffe und Schwierigkeitsprofile sind vorhanden.", "Boss-Grösse, Bodenkontakt und visuellen Schwerpunkt im Spiel prüfen."),
    ("Witch Boss", "EINGEBUNDEN / QA", "Boss-Entity, Renderer, Tripo-Mesh, 4K-Textur, Hasenform und Jagdphase sind vorhanden.", "Bossfight und Hasen-Verwandlung vollständig im Spiel testen."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_cell(cell, bold: bool = False, color: str = "1F2937", center: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(9)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)


def main() -> None:
    document = Document(DOCX)
    if any(paragraph.text.strip() == SECTION_TITLE for paragraph in document.paragraphs):
        raise SystemExit(f"Section already exists: {SECTION_TITLE}")

    heading = document.add_heading(SECTION_TITLE, level=1)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)

    note = document.add_paragraph(
        "Dieser Nachtrag dokumentiert den tatsächlich eingebundenen Projektstand. "
        "Status ‚QA‘ bedeutet: Code und Laufzeitressourcen sind vorhanden, die abschliessende isolierte Ingame-Sichtprüfung ist noch offen."
    )
    note.paragraph_format.space_after = Pt(8)

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.1), Inches(1.0), Inches(2.3), Inches(2.1)]
    headers = ["Mob", "Status", "Aktueller Stand", "Nächste Qualitätsprüfung"]
    for index, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = width
        cell.text = header
        set_cell_shading(cell, "273449")
        style_cell(cell, bold=True, color="FFFFFF", center=index < 2)
    set_repeat_header(table.rows[0])

    for row_index, values in enumerate(ROWS, start=1):
        cells = table.add_row().cells
        for index, (cell, value, width) in enumerate(zip(cells, values, widths)):
            cell.width = width
            cell.text = value
            if row_index % 2 == 0:
                set_cell_shading(cell, "EEF3F8")
            style_cell(cell, bold=index == 0, center=index == 1)

    document.save(DOCX)


if __name__ == "__main__":
    main()
